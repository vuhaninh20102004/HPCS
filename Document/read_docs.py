import sys, subprocess
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx

def convert_to_md(path, out_path):
    try:
        doc = docx.Document(path)
        with open(out_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                if para.text.strip():
                    f.write(para.text + '\n')
            f.write('\n\n--- TABLES ---\n\n')
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                    f.write(' | '.join(row_data) + '\n')
                f.write('\n')
    except Exception as e:
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(str(e))

if __name__ == '__main__':
    convert_to_md(r'e:\Capstone 1\HPCS\Document\C1SE.08_ProductBacklog_HPCS_ver1.0.docx', r'e:\Capstone 1\HPCS\Document\pb.md')
    convert_to_md(r'e:\Capstone 1\HPCS\Document\C1SE.08_UserStory_HPCS_ver1.0.docx', r'e:\Capstone 1\HPCS\Document\us.md')
